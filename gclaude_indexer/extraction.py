# GClaude Indexer — document collection indexer
# Copyright (C) 2026  Alex Camacho Castilho
#
# This program is free software: you can redistribute it and/or modify it
# under the terms of the GNU General Public License as published by the Free
# Software Foundation, either version 3 of the License, or (at your option)
# any later version. See the LICENSE file for details.

"""Per-page extraction (section 5, step 4).

For each already-converted file, reads the text of each page (from the
final PDF — original or OCR'd — or from the `.txt` extracted in phase 3 for
the other formats), writes it to `page` with the text truncated at the
configured limit, character count, image count, and table detection.

Each page's citable reference follows the collection type:
- `processo`: "f. N", with N counting the document's real numbering — a
  running sum of pages across all files of the same group, in the order
  they appear (natural ordering by relative path).
- `biblioteca`: "p. N", restarting at 1 for each file.

Resumable: a file with `status = "extracted"` already has its pages written
and is skipped, but still counts toward the group's page tally so that the
following files keep numbering correctly.
"""

from __future__ import annotations

import re
from collections import defaultdict
from collections.abc import Callable
from concurrent.futures import CancelledError, ProcessPoolExecutor, as_completed
from concurrent.futures.process import BrokenProcessPool
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF

from .config import ProjectConfig
from .events import record_event
from .parallelism import workers_for
from .file_types import category_of_extension

ERROR_MESSAGE_CHAR_LIMIT = 500


@dataclass
class ExtractionResult:
    files_processed: int = 0
    pages_written: int = 0
    failed: int = 0


def _natural_sort_key(text: str):
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", text)]


def _read_pdf_pages(pdf_path: Path, limit: int) -> list[tuple[str, int, bool]]:
    document = fitz.open(pdf_path)
    try:
        pages = []
        for page in document:
            text = page.get_text()[:limit]
            image_count = len(page.get_images(full=True))
            try:
                has_table = len(page.find_tables().tables) > 0
            except Exception:
                has_table = False
            pages.append((text, image_count, has_table))
        return pages
    finally:
        document.close()


def _docx_has_table(path: Path) -> bool:
    from docx import Document

    return len(Document(path).tables) > 0


def _extract_file_pages(config: ProjectConfig, row) -> list[tuple[str, int, bool]]:
    # Truncate here, in the worker, not only in `_write_pages`: on the
    # parallel path the return value crosses `pickle` to the main process,
    # and without truncating first the **full** text of every page stayed in
    # RAM (and was serialized whole) until it was written — for a large
    # collection that's the entire collection's text in memory at once (I2,
    # Phase 13 final review). The sequential path never had this problem
    # (one file at a time), but truncates here too, at no extra cost.
    limit = config.chars_per_page
    category = category_of_extension("." + row["extension"])
    relative_path = row["relative_path"]
    output_dir = Path(config.output_folder)

    if category == "pdf":
        if row["needs_ocr"]:
            pdf_path = output_dir / "converted" / Path(relative_path).with_suffix(".pdf")
        else:
            pdf_path = Path(config.source_folder) / relative_path
        return _read_pdf_pages(pdf_path, limit)

    txt_path = output_dir / "converted" / Path(relative_path).with_suffix(".txt")
    text = txt_path.read_text(encoding="utf-8")[:limit]

    if category == "imagens":
        return [(text, 1, False)]
    if category == "xlsx":
        return [(text, 0, True)]
    if category == "docx":
        has_table = _docx_has_table(Path(config.source_folder) / relative_path)
        return [(text, 0, has_table)]
    return [(text, 0, False)]  # pptx, texto, email


def _write_pages(conn, config: ProjectConfig, row, pages: list[tuple[str, int, bool]], start_sheet: int) -> int:
    limit = config.chars_per_page
    library = config.collection_type == "biblioteca"
    sheet = start_sheet

    for index, (full_text, image_count, has_table) in enumerate(pages, start=1):
        sheet += 1
        truncated_text = full_text[:limit]
        reference = f"p. {index}" if library else f"f. {sheet}"

        conn.execute(
            """
            INSERT INTO page
                (file_id, number, reference, char_count, image_count, has_table, text)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (row["id"], index, reference, len(truncated_text), image_count, int(has_table), truncated_text),
        )

    return sheet


def _build_groups(conn) -> dict[str, list]:
    rows = conn.execute(
        "SELECT * FROM file WHERE status IN ('converted', 'extracted') ORDER BY relative_path"
    ).fetchall()

    groups: dict[str, list] = defaultdict(list)
    for row in rows:
        group_key = row["group_key"] or f"__no_group__{row['id']}"
        groups[group_key].append(row)

    for group_files in groups.values():
        group_files.sort(key=lambda r: _natural_sort_key(r["relative_path"]))

    return groups


def _extract_sequential(
    conn, config: ProjectConfig, groups: dict[str, list], should_stop, result, language: str | None = None
) -> bool:
    interrupted = False

    for group_files in groups.values():
        if interrupted:
            break

        running_sheet = 0

        for row in group_files:
            if should_stop is not None and should_stop():
                interrupted = True
                break

            if row["status"] == "extracted":
                running_sheet += row["page_count"] or 0
                continue

            try:
                pages = _extract_file_pages(config, row)
            except Exception as exc:
                error_message = str(exc)[:ERROR_MESSAGE_CHAR_LIMIT]
                conn.execute(
                    "UPDATE file SET status = 'failed', error = ? WHERE id = ?",
                    (error_message, row["id"]),
                )
                conn.commit()
                result.failed += 1
                record_event(
                    conn, "extraction", "error", "log.extraction.failed",
                    {"name": row["name"], "error": error_message}, language=language,
                )
                continue

            running_sheet = _write_pages(conn, config, row, pages, running_sheet)
            conn.execute("UPDATE file SET status = 'extracted' WHERE id = ?", (row["id"],))
            conn.commit()

            result.files_processed += 1
            result.pages_written += len(pages)

    return interrupted


def _extract_in_parallel(
    conn, config: ProjectConfig, groups: dict[str, list], workers: int, should_stop, result,
    language: str | None = None,
) -> bool:
    """Page reading (PDF/tables — the CPU/IO-heavy part) runs in separate
    processes, one file per worker. Sheet numbering ("f. N") depends on file
    order within each group, so writing a file can only happen after every
    earlier file in the same group has been written.

    This doesn't require waiting for the whole collection to finish
    computing before writing anything (I1, Phase 13 final review): we keep,
    per group, a "next position to apply" pointer and a buffer of results
    already done out of order. Each result is applied (written and
    committed) as soon as its turn comes — the progress bar, which counts
    `status = 'extracted'`, starts moving during computation, not only at
    the end.

    `sqlite3.Connection` never crosses a process boundary: workers only read
    files from disk and return `list[tuple[str, int, bool]]`; the one that
    always writes to the database is this main process.
    """
    current_position: dict[str, int] = {key: 0 for key in groups}
    current_sheet: dict[str, int] = {key: 0 for key in groups}
    ready: dict[str, dict[int, tuple[str, object]]] = {key: {} for key in groups}

    def _apply_ready(group_key: str) -> None:
        group_files = groups[group_key]
        while current_position[group_key] < len(group_files):
            index = current_position[group_key]
            row = group_files[index]

            if row["status"] == "extracted":
                current_sheet[group_key] += row["page_count"] or 0
                current_position[group_key] += 1
                continue

            item = ready[group_key].pop(index, None)
            if item is None:
                return  # not here yet — retry when the next future completes

            kind, value = item
            if kind == "error":
                error_message = str(value)[:ERROR_MESSAGE_CHAR_LIMIT]
                conn.execute(
                    "UPDATE file SET status = 'failed', error = ? WHERE id = ?",
                    (error_message, row["id"]),
                )
                conn.commit()
                result.failed += 1
                record_event(
                    conn, "extraction", "error", "log.extraction.failed",
                    {"name": row["name"], "error": error_message}, language=language,
                )
            else:
                pages = value
                current_sheet[group_key] = _write_pages(conn, config, row, pages, current_sheet[group_key])
                conn.execute("UPDATE file SET status = 'extracted' WHERE id = ?", (row["id"],))
                conn.commit()
                result.files_processed += 1
                result.pages_written += len(pages)

            current_position[group_key] += 1

    pending = [
        (group_key, index, row)
        for group_key, group_files in groups.items()
        for index, row in enumerate(group_files)
        if row["status"] != "extracted"
    ]

    with ProcessPoolExecutor(max_workers=workers) as executor:
        futures: dict = {}
        for group_key, index, row in pending:
            if should_stop is not None and should_stop():
                break
            future = executor.submit(_extract_file_pages, config, dict(row))
            futures[future] = (group_key, index, row)

        stop_request_handled = False
        pool_broken = False
        for future in as_completed(futures):
            group_key, index, row = futures[future]
            try:
                ready[group_key][index] = ("ok", future.result())
            except CancelledError:
                continue
            except BrokenProcessPool:
                # A worker died (e.g. `MemoryError` on a large PDF) and took
                # the whole pool down with it — ALL still-pending futures
                # receive this same exception, including files that never
                # got to run (I3, Phase 13 final review). Unlike a real
                # per-file failure: we don't mark it 'failed' (which would
                # exclude the file from every future rerun) — we leave this
                # item without a result in `ready`, which stalls the group's
                # `current_position` right here and keeps the original
                # status ('converted') to be resumed later.
                if not pool_broken:
                    pool_broken = True
                    record_event(conn, "extraction", "error", "log.pool_broken", language=language)
                continue
            except Exception as exc:
                ready[group_key][index] = ("error", exc)

            _apply_ready(group_key)

            if not stop_request_handled and should_stop is not None and should_stop():
                stop_request_handled = True
                # The pause button must never go unanswered: cancel every
                # future not yet started.
                for pending_future in futures:
                    if not pending_future.done():
                        pending_future.cancel()

    return any(current_position[key] < len(groups[key]) for key in groups)


def extract_pages(
    conn, config: ProjectConfig, should_stop: Callable[[], bool] | None = None, language: str | None = None
) -> ExtractionResult:
    groups = _build_groups(conn)
    result = ExtractionResult()

    workers = workers_for(config.parallelism)
    n_pending = sum(
        1 for group_files in groups.values() for row in group_files if row["status"] != "extracted"
    )
    # With 0 or 1 file pending there's nothing to parallelize — the
    # `ProcessPoolExecutor` startup cost (spawn, on Windows) outweighs any
    # gain.
    if workers <= 1 or n_pending <= 1:
        interrupted = _extract_sequential(conn, config, groups, should_stop, result, language=language)
    else:
        interrupted = _extract_in_parallel(conn, config, groups, workers, should_stop, result, language=language)

    summary_key = "log.extraction.summary_interrupted" if interrupted else "log.extraction.summary_completed"
    record_event(
        conn,
        "extraction",
        "info",
        summary_key,
        {"files": result.files_processed, "pages": result.pages_written, "failed": result.failed},
        language=language,
    )

    return result
