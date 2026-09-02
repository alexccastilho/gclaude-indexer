"""Phase 4 tests: page text extraction and window preparation."""

from __future__ import annotations

import fitz
from docx import Document

from gclaude_indexer.config import load_config
from gclaude_indexer.conversion import convert
from gclaude_indexer.db import connect, init_schema
from gclaude_indexer.extraction import extract_pages
from gclaude_indexer.windows_prep import generate_claude_md, prepare_windows
from gclaude_indexer.scanning import scan

TEXTO_LONGO = (
    "Texto nativo de teste, com camada de texto real e conteúdo suficiente "
    "para superar o limiar médio de cem caracteres por página."
)


def _criar_pdf(caminho, n_paginas=1, texto=TEXTO_LONGO):
    documento = fitz.open()
    for indice in range(n_paginas):
        pagina = documento.new_page()
        pagina.insert_textbox((50, 50, 550, 750), f"{texto} (página {indice + 1})", fontsize=12)
    documento.save(caminho)
    documento.close()


def _preparar_projeto(tmp_path, tipo_acervo="processo", extra=None):
    origem = tmp_path / "origem"
    origem.mkdir()
    saida = tmp_path / "origem_indexado"

    dados = {
        "name": "Projeto Fase 4",
        "subject": "Acervo de teste da fase 4",
        "source_folder": str(origem),
        "output_folder": str(saida),
        "collection_type": tipo_acervo,
    }
    if extra:
        dados.update(extra)
    config = load_config(dados)

    conn = connect(saida / "project.db")
    init_schema(conn)
    return origem, saida, config, conn


def _rodar_ate_extracao(origem, config, conn):
    scan(conn, config)
    convert(conn, config)
    return extract_pages(conn, config)


# --- referência por tipo de acervo -----------------------------------------


def test_processo_numera_folha_continua_entre_arquivos_do_mesmo_agrupador(tmp_path):
    origem, saida, config, conn = _preparar_projeto(tmp_path, tipo_acervo="processo")
    (origem / "volume_1").mkdir()
    _criar_pdf(origem / "volume_1" / "peca_a.pdf", n_paginas=2)
    _criar_pdf(origem / "volume_1" / "peca_b.pdf", n_paginas=3)

    _rodar_ate_extracao(origem, config, conn)

    linhas = conn.execute(
        """
        SELECT file.name AS arquivo_nome, page.number, page.reference
        FROM page JOIN file ON file.id = page.file_id
        ORDER BY page.id
        """
    ).fetchall()

    referencias = [(l["arquivo_nome"], l["number"], l["reference"]) for l in linhas]
    assert referencias == [
        ("peca_a.pdf", 1, "f. 1"),
        ("peca_a.pdf", 2, "f. 2"),
        ("peca_b.pdf", 1, "f. 3"),
        ("peca_b.pdf", 2, "f. 4"),
        ("peca_b.pdf", 3, "f. 5"),
    ]

    conn.close()


def test_biblioteca_reinicia_pagina_a_cada_arquivo(tmp_path):
    origem, saida, config, conn = _preparar_projeto(tmp_path, tipo_acervo="biblioteca")
    (origem / "volume_1").mkdir()
    _criar_pdf(origem / "volume_1" / "capitulo1.pdf", n_paginas=2)
    _criar_pdf(origem / "volume_1" / "capitulo2.pdf", n_paginas=2)

    _rodar_ate_extracao(origem, config, conn)

    linhas = conn.execute(
        """
        SELECT file.name AS arquivo_nome, page.reference
        FROM page JOIN file ON file.id = page.file_id
        ORDER BY page.id
        """
    ).fetchall()

    referencias = [(l["arquivo_nome"], l["reference"]) for l in linhas]
    assert referencias == [
        ("capitulo1.pdf", "p. 1"),
        ("capitulo1.pdf", "p. 2"),
        ("capitulo2.pdf", "p. 1"),
        ("capitulo2.pdf", "p. 2"),
    ]

    conn.close()


def test_ordenacao_natural_evita_ue10_antes_de_ue2(tmp_path):
    origem, saida, config, conn = _preparar_projeto(tmp_path, tipo_acervo="processo")
    (origem / "curso").mkdir()
    _criar_pdf(origem / "curso" / "UE1.pdf", n_paginas=1)
    _criar_pdf(origem / "curso" / "UE2.pdf", n_paginas=1)
    _criar_pdf(origem / "curso" / "UE10.pdf", n_paginas=1)

    _rodar_ate_extracao(origem, config, conn)

    ordem = [
        row["nome"]
        for row in conn.execute(
            """
            SELECT file.name AS nome
            FROM page JOIN file ON file.id = page.file_id
            ORDER BY page.id
            """
        ).fetchall()
    ]
    assert ordem == ["UE1.pdf", "UE2.pdf", "UE10.pdf"]

    conn.close()


# --- truncamento e sinais estruturais -----------------------------------


def test_texto_truncado_no_limite_configurado(tmp_path):
    origem, saida, config, conn = _preparar_projeto(
        tmp_path, extra={"chars_per_page": 20}
    )
    _criar_pdf(origem / "peca.pdf", n_paginas=1)

    _rodar_ate_extracao(origem, config, conn)

    linha = conn.execute("SELECT * FROM page").fetchone()
    assert len(linha["text"]) <= 20
    assert linha["char_count"] == len(linha["text"])

    conn.close()


def test_docx_marca_tem_tabela_quando_ha_tabela(tmp_path):
    origem, saida, config, conn = _preparar_projeto(tmp_path)
    documento = Document()
    documento.add_paragraph("Texto antes da tabela.")
    tabela = documento.add_table(rows=2, cols=2)
    tabela.cell(0, 0).text = "a"
    documento.save(origem / "com_tabela.docx")

    _rodar_ate_extracao(origem, config, conn)

    linha = conn.execute("SELECT * FROM page").fetchone()
    assert linha["has_table"] == 1
    assert linha["image_count"] == 0

    conn.close()


# --- retomabilidade -----------------------------------------------------


def test_extracao_e_retomavel_sem_duplicar_paginas(tmp_path):
    origem, saida, config, conn = _preparar_projeto(tmp_path)
    (origem / "volume_1").mkdir()
    _criar_pdf(origem / "volume_1" / "peca_a.pdf", n_paginas=2)

    _rodar_ate_extracao(origem, config, conn)
    total_apos_primeira = conn.execute("SELECT COUNT(*) FROM page").fetchone()[0]

    resultado_segunda = extract_pages(conn, config)
    assert resultado_segunda.files_processed == 0
    assert conn.execute("SELECT COUNT(*) FROM page").fetchone()[0] == total_apos_primeira

    # arquivo novo no mesmo agrupador continua a numeração de folha
    _criar_pdf(origem / "volume_1" / "peca_b.pdf", n_paginas=1)
    scan(conn, config)
    convert(conn, config)
    extract_pages(conn, config)

    nova = conn.execute(
        """
        SELECT page.reference FROM page
        JOIN file ON file.id = page.file_id
        WHERE file.name = 'peca_b.pdf'
        """
    ).fetchone()
    assert nova["reference"] == "f. 3"

    conn.close()


# --- janelas -------------------------------------------------------------


def test_preparar_janelas_gera_arquivos_com_sobreposicao(tmp_path):
    origem, saida, config, conn = _preparar_projeto(
        tmp_path, extra={"pages_per_window": 3, "overlap": 1}
    )
    (origem / "volume_1").mkdir()
    _criar_pdf(origem / "volume_1" / "peca.pdf", n_paginas=5)

    _rodar_ate_extracao(origem, config, conn)
    resultado = prepare_windows(conn, config)

    assert resultado.created == 2  # [f.1-f.3], [f.3-f.5] (passo = 3-1 = 2)

    janelas = conn.execute("SELECT * FROM window ORDER BY id").fetchall()
    assert [j["start_ref"] for j in janelas] == ["f. 1", "f. 3"]
    assert [j["end_ref"] for j in janelas] == ["f. 3", "f. 5"]
    assert all(j["status"] == "pending" for j in janelas)

    arquivos_txt = sorted((saida / "windows").glob("*.txt"))
    assert len(arquivos_txt) == 2

    conteudo_primeira = arquivos_txt[0].read_text(encoding="utf-8")
    assert "volume_1" in conteudo_primeira
    assert "f. 1" in conteudo_primeira
    assert "f. 3" in conteudo_primeira
    assert "página 1" in conteudo_primeira  # texto real da página incluído

    conn.close()


def test_preparar_janelas_e_idempotente(tmp_path):
    origem, saida, config, conn = _preparar_projeto(
        tmp_path, extra={"pages_per_window": 3, "overlap": 1}
    )
    (origem / "volume_1").mkdir()
    _criar_pdf(origem / "volume_1" / "peca.pdf", n_paginas=5)

    _rodar_ate_extracao(origem, config, conn)
    prepare_windows(conn, config)
    total_apos_primeira = conn.execute("SELECT COUNT(*) FROM window").fetchone()[0]

    resultado_segunda = prepare_windows(conn, config)
    assert resultado_segunda.created == 0
    assert resultado_segunda.existing == 2
    assert conn.execute("SELECT COUNT(*) FROM window").fetchone()[0] == total_apos_primeira

    conn.close()


# --- CLAUDE.md -----------------------------------------------------------


def test_gerar_claude_md_contem_loop_e_formato_json(tmp_path):
    origem, saida, config, conn = _preparar_projeto(tmp_path)

    caminho = generate_claude_md(config, "pt")

    assert caminho == saida / "CLAUDE.md"
    conteudo = caminho.read_text(encoding="utf-8")
    assert "processe as janelas" in conteudo
    assert "raw_items.jsonl" in conteudo
    assert '"engine": "claude_code"' in conteudo
    assert '"confidence"' in conteudo
    assert "order_start" in conteudo and "order_end" in conteudo
    assert config.subject in conteudo

    conn.close()
